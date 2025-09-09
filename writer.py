import os, re, sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import config

# Caractères invisibles qu’on rencontre souvent en copiant-collant du texte
NBSP = "\u00A0"
ZWS  = "\u200B"

# On repère une éventuelle salutation au tout début (FR/EN) pour pouvoir la retirer proprement
GREET_RX = re.compile(
    r"^\s*(dear(\s+hiring\s+(team|manager))?|bonjour|madame|monsieur|a\s+l'attention|à\s+l'attention)\b[^\n]*?,?\s*",
    re.IGNORECASE,
)
# On repère les formules de fin de lettre (FR/EN) pour les ignorer si l’utilisateur les a déjà mises
CLOSE_RX = re.compile(
    r"\b(yours\s+sincerely|kind\s+regards|best\s+regards|cordialement)\b.*$",
    re.IGNORECASE,
)

def app_dir() -> str:
    # Où sauvegarder les fichiers :
    # - si l’app est “gelée” (exécutable), on prend le dossier de l’exe
    # - sinon, le dossier du script Python
    if getattr(sys, "frozen", False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))

# ————— Utilitaires (petites briques réutilisables pour garder le code lisible) —————
def _cfg(name, default):
    # Récupère une valeur depuis config.py, avec un secours si la clé n’existe pas
    return getattr(config, name, default)

def _normalize_ws(s: str) -> str:
    # Nettoie les espaces chelous (NBSP/ZWS), enlève puces/chevrons en début de ligne, et trim
    s = (s or "").replace(NBSP, " ").replace(ZWS, " ").strip()
    s = re.sub(r"^[\u2022>\-\*\•\s]+", "", s)
    return s

def _apply_font_to_run(run, bold=False):
    """Applique notre police/taille (y compris au niveau XML de Word) et met en gras si demandé."""
    run.bold = bool(bold)
    run.font.name = _cfg("POLICE", "Aptos")
    run.font.size = Pt(_cfg("TAILLE_PT", 12))

    # Ci-dessous : on ajuste aussi les propriétés au niveau XML pour que Word soit 100% d’accord
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr"); r.append(rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rPr.append(rFonts)
    font_name = _cfg("POLICE", "Aptos")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)

def _add_text(p, text, bold=False):
    # Ajoute un run de texte au paragraphe avec le style maison
    r = p.add_run(text)
    _apply_font_to_run(r, bold=bold)
    return r

def _mailto(paragraph, email: str):
    # Crée un lien cliquable “mailto:” dans le document (style Word natif “Hyperlink”)
    part = paragraph.part
    r_id = part.relate_to(f"mailto:{email}", RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink"); hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle"); rStyle.set(qn("w:val"), "Hyperlink"); rPr.append(rStyle)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)

    # On garde notre police/taille même sur le lien pour l’uniformité visuelle
    rFonts = OxmlElement("w:rFonts")
    font_name = _cfg("POLICE", "Aptos")
    rFonts.set(qn("w:ascii"), font_name); rFonts.set(qn("w:hAnsi"), font_name); rFonts.set(qn("w:cs"), font_name)
    rPr.append(rFonts)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(_cfg("TAILLE_PT", 12) * 2)))
    szCs = OxmlElement("w:szCs"); szCs.set(qn("w:val"), str(int(_cfg("TAILLE_PT", 12) * 2)))
    rPr.append(sz); rPr.append(szCs)
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rPr.append(color)

    run.append(rPr)
    t = OxmlElement("w:t"); t.text = email; run.append(t)
    hyperlink.append(run); paragraph._p.append(hyperlink)

def _clean(text: str) -> str:
    # Nettoyage léger : espaces, markdown basique (* _ `), doublons d’espaces
    s = _normalize_ws(text)
    s = re.sub(r"[\*`_]+", "", s)
    s = re.sub(r"[ \t]{2,}", " ", s)
    return s.strip()

def _pf(p, after_pt=0, before_pt=0):
    # Applique notre mise en forme de paragraphe (interligne + espacements avant/après)
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = _cfg("LINE_SPACING", 1.22)
    pf.space_before = Pt(before_pt)
    pf.space_after  = Pt(after_pt)

def _sanitize_paragraphs(pars: list[str]) -> list[str]:
    # Prépare proprement les paragraphes fournis par l’utilisateur avant de les injecter dans Word
    out = []
    for para in pars or []:
        s = _clean(para)
        # Si l’utilisateur a déjà commencé par “Bonjour/Dear…”, on enlève la salutation (on garde le reste)
        s = GREET_RX.sub("", s)
        # Si le paragraphe est en fait une formule de politesse de fin, on la laissera plutôt à la signature
        if CLOSE_RX.search(s):
            continue
        if s:
            out.append(s)
    # On se limite à 4 paragraphes pour garder la lettre concise et lisible
    return out[:4]

# ————— Cœur du sujet : construire la lettre et l’enregistrer —————
def build_letter_doc(bank: str, position: str, body_paragraphs: list[str]) -> Document:
    body_paragraphs = _sanitize_paragraphs(body_paragraphs)

    doc = Document()

    # Marges du document (tirées de config, avec des valeurs par défaut raisonnables)
    top, bottom, left, right = _cfg("MARGES_INCH", (1.00, 0.60, 1.10, 1.10))
    for s in doc.sections:
        s.top_margin = Inches(top)
        s.bottom_margin = Inches(bottom)
        s.left_margin  = Inches(left)
        s.right_margin = Inches(right)

    # Style “Normal” de Word : on l’aligne avec nos choix (police, taille, interligne, pas d’espaces parasites)
    st = doc.styles["Normal"]
    st.font.name = _cfg("POLICE", "Aptos")
    st.font.size = Pt(_cfg("TAILLE_PT", 12))
    st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    st.paragraph_format.line_spacing = _cfg("LINE_SPACING", 1.22)
    st.paragraph_format.space_before = Pt(0)
    st.paragraph_format.space_after  = Pt(0)

    # Espacements verticaux entre les blocs (tous configurables)
    after_email     = _cfg("ESP_APRES_EMAIL_PT",     12)
    after_subject   = _cfg("ESP_APRES_SUJET_PT",     12)
    after_salut     = _cfg("ESP_APRES_SALUT_PT",     12)
    after_par       = _cfg("ESP_APRES_PAR_PT",         6)
    after_signature = _cfg("ESP_APRES_SIGNATURE_PT", 12)

    # Bloc identité en en-tête (nom, adresse, ville, téléphone)
    p = doc.add_paragraph()
    _add_text(p, config.NOM, bold=True); _add_text(p, "\n")
    _add_text(p, config.ADRESSE); _add_text(p, "\n")
    _add_text(p, config.VILLE_PAYS); _add_text(p, "\n")
    _add_text(p, config.TEL)
    _pf(p, after_pt=0)

    # Email cliquable (mailto:) juste en dessous
    p_mail = doc.add_paragraph(); _mailto(p_mail, config.EMAIL)
    _pf(p_mail, after_pt=after_email)

    # Sujet explicite avec le poste ciblé (utile côté RH et ATS)
    p_subj = doc.add_paragraph()
    _add_text(p_subj, f"Subject: Application – {position}", bold=True)
    _pf(p_subj, after_pt=after_subject)

    # Salutation : on l’insère une seule fois, le corps ayant été nettoyé de toute salutation redondante
    p_greet = doc.add_paragraph()
    _add_text(p_greet, "Dear Hiring Team,")
    _pf(p_greet, after_pt=after_salut)

    # Corps de la lettre (jusqu’à 4 paragraphes propres et aérés)
    for para in body_paragraphs[:4]:
        pb = doc.add_paragraph()
        _add_text(pb, _clean(para))
        _pf(pb, after_pt=after_par)

    # Formule de politesse + signature (on se charge de la cohérence stylistique ici)
    ps1 = doc.add_paragraph(); _add_text(ps1, "Yours sincerely,")
    _pf(ps1, after_pt=after_signature)
    ps2 = doc.add_paragraph(); _add_text(ps2, config.NOM, bold=True)
    _pf(ps2, after_pt=0)

    return doc

def save_letter(bank: str, position: str, body_paragraphs: list[str]) -> str:
    # Construit le document puis l’écrit sur disque dans un dossier par banque
    doc = build_letter_doc(bank, position, body_paragraphs)
    out_root = os.path.join(app_dir(), _cfg("OUT_DIR", "generated_letters"))
    out_dir = os.path.join(out_root, bank)
    os.makedirs(out_dir, exist_ok=True)
    filename = f"Cover Letter {config.NOM} - {bank} - {position.replace(' ', '_')}.docx"
    path = os.path.join(out_dir, filename)

    # Si le fichier est déjà ouvert par Word (verrou), on suffixe automatiquement (1), (2), …
    base, ext = os.path.splitext(path); n = 1
    while True:
        try:
            doc.save(path); break
        except PermissionError:
            path = f"{base} ({n}){ext}"; n += 1
    return path
